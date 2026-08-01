"""Tests for the /contract/upload endpoint + the upload_extract
parser that backs it.

The endpoint is exercised via FastAPI's TestClient (no real
server needed). The parser unit tests cover each of the 4
supported formats (txt/md/pdf/docx/html) using synthetic
test files.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from dpo_agent.examples.fastapi_server import MAX_UPLOAD_BYTES, app
from dpo_agent.upload_extract import (
    ExtractionError,
    SUPPORTED_EXTENSIONS,
    UnsupportedFormatError,
    extract_text,
)


# Reuse the same TestFile fixtures from the existing Playwright
# test file location. Both upload tests (the Playwright ones
# and these FastAPI ones) need real PDF/DOCX content.
@pytest.fixture
def pdf_path(tmp_path: Path) -> Path:
    """Build a minimal 2-page PDF for tests."""
    from reportlab.pdfgen import canvas
    p = tmp_path / "tiny.pdf"
    c = canvas.Canvas(str(p))
    c.drawString(100, 750, "Master Services Agreement")
    c.drawString(100, 720, "Between Acme Corp (Provider) and Widget Inc (Customer).")
    c.drawString(100, 690, "Effective 2024-01-15.")
    c.showPage()
    c.drawString(100, 750, "Data Protection: Provider shall implement TOMs.")
    c.save()
    return p


@pytest.fixture
def docx_path(tmp_path: Path) -> Path:
    """Build a minimal DOCX with two sections."""
    from docx import Document
    p = tmp_path / "tiny.docx"
    doc = Document()
    doc.add_heading("Master Services Agreement", level=1)
    doc.add_paragraph("Between Acme Corp (Provider) and Widget Inc (Customer).")
    doc.add_heading("Data Protection", level=2)
    doc.add_paragraph("Provider shall implement appropriate TOMs.")
    doc.save(str(p))
    return p


# ── Parser unit tests ──────────────────────────────────


def test_extract_text_plain_utf8():
    """TXT files are decoded as UTF-8."""
    text = extract_text("Hello world\n".encode(), "test.txt")
    assert text == "Hello world\n"


def test_extract_text_markdown_via_dotted_ext():
    """Both .md and .markdown extensions are accepted."""
    for ext in (".md", ".markdown"):
        text = extract_text(b"# Heading\n\nBody", f"file{ext}")
        assert "# Heading" in text


def test_extract_text_pdf(pdf_path):
    """PDFs are parsed via pdfplumber; section_path appears as
    'Page N' in the extracted output."""
    text = extract_text(pdf_path.read_bytes(), "contract.pdf")
    assert "Master Services Agreement" in text
    assert "Acme Corp" in text
    # Page markers are injected for downstream LLM visibility.
    assert "Page 1" in text
    assert "Page 2" in text


def test_extract_text_docx(docx_path):
    """DOCX are parsed via python-docx."""
    text = extract_text(docx_path.read_bytes(), "contract.docx")
    assert "Master Services Agreement" in text
    assert "Data Protection" in text
    assert "TOMs" in text


def test_extract_text_rejects_unsupported_extension():
    """ZIP / unknown extensions raise UnsupportedFormatError."""
    with pytest.raises(UnsupportedFormatError) as exc_info:
        extract_text(b"PK\x03\x04 fake zip", "test.zip")
    assert "zip" in str(exc_info.value).lower()


def test_extract_text_rejects_oversize_content():
    """Bytes exceeding max_bytes are rejected with ValueError."""
    with pytest.raises(ValueError, match="too large"):
        extract_text(b"x" * 100, "test.txt", max_bytes=10)


def test_extract_text_corrupt_pdf():
    """A PDF with invalid bytes raises ExtractionError, not
    a parse traceback."""
    with pytest.raises(ExtractionError):
        extract_text(b"%PDF-1.4\nnot actually a pdf\n", "fake.pdf")


# ── Endpoint behavior tests ─────────────────────────────


@pytest.fixture
def client():
    return TestClient(app)


def test_upload_endpoint_accepts_pdf(pdf_path, client):
    """A valid PDF upload returns 200 with extracted text."""
    r = client.post(
        "/contract/upload",
        files={"file": (
            "contract.pdf",
            pdf_path.read_bytes(),
            "application/pdf",
        )},
    )
    assert r.status_code == 200
    data = r.json()
    assert data["filename"] == "contract.pdf"
    assert data["format"] == "pdf"
    assert "Master Services Agreement" in data["text"]
    assert "char_count" in data
    assert data["char_count"] > 0
    assert data["size"] > 0


def test_upload_endpoint_accepts_docx(docx_path, client):
    """A valid DOCX upload returns 200."""
    r = client.post(
        "/contract/upload",
        files={"file": (
            "contract.docx",
            docx_path.read_bytes(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # noqa: E501
        )},
    )
    assert r.status_code == 200
    data = r.json()
    assert data["filename"] == "contract.docx"
    assert data["format"] == "docx"
    assert "Master Services Agreement" in data["text"]


def test_upload_endpoint_accepts_markdown(client):
    """MD files (which would normally go through the
    client-side FileReader) are also accepted by the endpoint
    for flexibility — the response is the same shape."""
    r = client.post(
        "/contract/upload",
        files={"file": (
            "contract.md",
            b"# Master Services Agreement\n\nHello",
            "text/markdown",
        )},
    )
    assert r.status_code == 200
    data = r.json()
    assert data["format"] == "md"


def test_upload_endpoint_rejects_unsupported(client):
    """A .zip file gets a 400 with the supported-formats list."""
    r = client.post(
        "/contract/upload",
        files={"file": ("fake.zip", b"PK\x03\x04 fake", "application/zip")},
    )
    assert r.status_code == 400
    detail = r.json()["detail"]
    assert detail["error"] == "unsupported_format"
    assert ".pdf" in detail["supported_formats"]
    assert ".docx" in detail["supported_formats"]
    assert ".txt" in detail["supported_formats"]


def test_upload_endpoint_rejects_missing_filename(client):
    """An empty filename gets rejected. FastAPI's own validation
    returns 422 before our handler sees it, which is fine —
    the contract is "not 200 and the file isn't processed"."""
    r = client.post(
        "/contract/upload",
        files={"file": ("", b"content", "text/plain")},
    )
    assert r.status_code in (400, 422)
    # Body should not contain a successful extraction.
    if r.status_code == 400:
        assert "filename" in r.json()["detail"].lower()


def test_upload_endpoint_rejects_oversize(client):
    """Files larger than MAX_UPLOAD_BYTES return 413."""
    big = b"x" * (MAX_UPLOAD_BYTES + 1)
    r = client.post(
        "/contract/upload",
        files={"file": ("huge.txt", big, "text/plain")},
    )
    assert r.status_code == 413


def test_upload_endpoint_handles_corrupt_pdf(client):
    """A PDF with garbage bytes returns 422 (extraction_failed)."""
    r = client.post(
        "/contract/upload",
        files={"file": (
            "corrupt.pdf",
            b"%PDF-1.4\nnot actually a pdf body\n",
            "application/pdf",
        )},
    )
    assert r.status_code == 422
    detail = r.json()["detail"]
    assert detail["error"] == "extraction_failed"
    assert detail["filename"] == "corrupt.pdf"


# ── Supported-formats contract (parity with frontend) ──


def test_supported_extensions_includes_pdf_docx_html():
    """The set of supported extensions must include the formats
    the frontend advertises (so the server doesn't 400 a file
    the user was told they could upload)."""
    assert ".pdf" in SUPPORTED_EXTENSIONS
    assert ".docx" in SUPPORTED_EXTENSIONS
    assert ".html" in SUPPORTED_EXTENSIONS
    assert ".htm" in SUPPORTED_EXTENSIONS
    assert ".md" in SUPPORTED_EXTENSIONS
    assert ".txt" in SUPPORTED_EXTENSIONS
